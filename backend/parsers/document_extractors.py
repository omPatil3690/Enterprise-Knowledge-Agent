"""
Enterprise Document Parsers for Binary File Formats.

Extracts structured semantic blocks (headings, paragraphs, bullet points, tabular databases)
from binary enterprise formats:
- PDF (.pdf) via pypdf
- Microsoft Word (.docx) via python-docx
- Microsoft Excel (.xlsx) via openpyxl
"""

import io
import re
from typing import Any, Dict, List, Optional
import pypdf
import docx
import openpyxl


def build_heading(text: str, level: int = 1) -> Dict[str, Any]:
    """Builds a normalized heading block dictionary."""
    clean = text.strip()
    return {
        "type": f"heading_{min(max(level, 1), 3)}",
        "text": clean,
        "block_id": re.sub(r"[^\w\s-]", "", clean).lower().replace(" ", "_")[:60] or "heading",
        "properties": {"level": level},
    }


def build_paragraph(text: str) -> Dict[str, Any]:
    """Builds a normalized paragraph block dictionary."""
    clean = text.strip()
    return {
        "type": "paragraph",
        "text": clean,
        "block_id": re.sub(r"[^\w\s-]", "", clean).lower().replace(" ", "_")[:60] or "para",
    }


def build_bullet(text: str) -> Dict[str, Any]:
    """Builds a normalized bulleted list item block dictionary."""
    clean = text.strip()
    return {
        "type": "bullet",
        "text": clean,
        "block_id": re.sub(r"[^\w\s-]", "", clean).lower().replace(" ", "_")[:60] or "bullet",
    }


def extract_pdf_blocks(file_bytes: bytes, filename: str = "") -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF document.
    
    Returns:
        List of ContentBlock dicts with page headings and extracted paragraphs.
    """
    blocks: List[Dict[str, Any]] = []
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        total_pages = len(reader.pages)

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue

            if total_pages > 1:
                blocks.append(build_heading(f"Page {page_num}", level=2))

            # Split paragraphs on double newlines
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            for p in paragraphs:
                blocks.append(build_paragraph(p))

    except Exception as e:
        blocks.append(build_paragraph(f"[PDF Extraction Error: {e}]"))

    return blocks


def extract_docx_blocks(file_bytes: bytes, filename: str = "") -> List[Dict[str, Any]]:
    """
    Extracts structured headings, paragraphs, bullet lists, and tables from a Word (.docx) document.
    
    Returns:
        List of ContentBlock dicts preserving document structure.
    """
    blocks: List[Dict[str, Any]] = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))

        # 1. Extract paragraphs & headings
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = (p.style.name or "").lower()
            if "heading 1" in style_name:
                blocks.append(build_heading(text, level=1))
            elif "heading 2" in style_name:
                blocks.append(build_heading(text, level=2))
            elif "heading 3" in style_name:
                blocks.append(build_heading(text, level=3))
            elif "heading" in style_name or "title" in style_name:
                blocks.append(build_heading(text, level=1))
            elif "list" in style_name or "bullet" in style_name:
                blocks.append(build_bullet(text))
            else:
                blocks.append(build_paragraph(text))

        # 2. Extract tables
        for t_idx, table in enumerate(doc.tables, 1):
            if not table.rows:
                continue

            # First row is typically header
            raw_headers = [cell.text.strip() for cell in table.rows[0].cells]
            # Deduplicate and ensure non-empty headers
            headers: List[str] = []
            for i, h in enumerate(raw_headers, 1):
                headers.append(h if h else f"Column_{i}")

            rows: List[Dict[str, Any]] = []
            for r_idx, row in enumerate(table.rows[1:], 2):
                row_data: Dict[str, Any] = {}
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(headers):
                        row_data[headers[col_idx]] = cell.text.strip()
                if any(row_data.values()):
                    rows.append({
                        "id": f"table_{t_idx}_row_{r_idx}",
                        "data": row_data,
                    })

            if rows:
                blocks.append({
                    "type": "database",
                    "text": f"Table {t_idx}",
                    "block_id": f"table_{t_idx}",
                    "columns": headers,
                    "rows": rows,
                    "total_rows": len(rows),
                })

    except Exception as e:
        blocks.append(build_paragraph(f"[DOCX Extraction Error: {e}]"))

    return blocks


def extract_xlsx_blocks(file_bytes: bytes, filename: str = "") -> List[Dict[str, Any]]:
    """
    Extracts sheets, columns, and rows from an Excel (.xlsx) workbook into structured database blocks.
    
    Returns:
        List of ContentBlock dicts representing spreadsheet tables.
    """
    blocks: List[Dict[str, Any]] = []
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            sheet_rows = list(ws.iter_rows(values_only=True))
            if not sheet_rows:
                continue

            # Find first non-empty row to use as headers
            header_row_idx = -1
            raw_headers: List[Any] = []
            for idx, r in enumerate(sheet_rows):
                if any(val is not None and str(val).strip() for val in r):
                    header_row_idx = idx
                    raw_headers = list(r)
                    break

            if header_row_idx == -1:
                continue

            headers: List[str] = []
            for c_idx, h in enumerate(raw_headers, 1):
                col_name = str(h).strip() if h is not None else ""
                headers.append(col_name if col_name else f"Column_{c_idx}")

            # Extract data rows
            data_rows: List[Dict[str, Any]] = []
            for r_idx, r in enumerate(sheet_rows[header_row_idx + 1:], header_row_idx + 2):
                row_data: Dict[str, Any] = {}
                for c_idx, cell_value in enumerate(r):
                    if c_idx < len(headers):
                        val_str = "" if cell_value is None else str(cell_value).strip()
                        row_data[headers[c_idx]] = val_str
                if any(row_data.values()):
                    data_rows.append({
                        "id": f"{sheet_name}_row_{r_idx}",
                        "data": row_data,
                    })

            if data_rows:
                blocks.append(build_heading(f"Sheet: {sheet_name}", level=2))
                blocks.append({
                    "type": "database",
                    "text": f"Worksheet: {sheet_name}",
                    "block_id": f"sheet_{sheet_name.lower().replace(' ', '_')}",
                    "columns": headers,
                    "rows": data_rows,
                    "total_rows": len(data_rows),
                })

        wb.close()

    except Exception as e:
        blocks.append(build_paragraph(f"[XLSX Extraction Error: {e}]"))

    return blocks


def extract_document_blocks(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Unified router that extracts semantic content blocks from any supported binary or text file.
    
    Supports:
    - .pdf  -> extract_pdf_blocks
    - .docx -> extract_docx_blocks
    - .xlsx -> extract_xlsx_blocks
    - .txt / .md / code -> decodes plain text into paragraphs
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return extract_pdf_blocks(file_bytes, filename)
    elif ext in ("docx", "doc"):
        return extract_docx_blocks(file_bytes, filename)
    elif ext in ("xlsx", "xls"):
        return extract_xlsx_blocks(file_bytes, filename)
    else:
        # Fallback to text decoding
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="replace")

        blocks: List[Dict[str, Any]] = []
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for p in paragraphs:
            blocks.append(build_paragraph(p))
        return blocks
